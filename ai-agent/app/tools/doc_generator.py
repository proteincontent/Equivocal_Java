from langchain_core.tools import tool
from docx import Document
from docx.oxml.ns import qn
from app.services.storage import storage_service
import io
import os
import tempfile

@tool
def generate_legal_document(title: str, content: str, doc_type: str = "contract") -> str:
    """Generate a Word document and return the download URL."""
    try:
        # Create a new Document
        doc = Document()
        
        # Set default font to SimSun (宋体)
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc.add_heading(title, 0)
        
        # Add content paragraphs
        # Simple splitting by newline for paragraphs
        for paragraph in content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
            
        # Read the file back to bytes
        with open(tmp_path, "rb") as f:
            file_data = f.read()
            
        # Clean up temp file
        os.unlink(tmp_path)
        
        # Upload to R2
        # 清理文件名：移除特殊字符，用下划线替换空格
        safe_title = title.replace(' ', '_').replace(':', '').replace('/', '_').replace('\\', '_')
        filename = f"{doc_type}_{safe_title}.docx"
        url = storage_service.upload_file(file_data, filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        # 返回中文提示和 Markdown 链接，确保前端渲染为可点击的文件下载
        return f"文档已生成完成。\n\n[📄 点击下载《{title}》]({url})\n\n提示：请仔细核对文档中的 [待定] 内容，根据实际情况填写完整。"
        
    except Exception as e:
        return f"文档生成失败：{str(e)}"

@tool
def list_supported_documents() -> str:
    """List supported legal document types."""
    return """
    Supported Document Types:
    1. 房屋租赁合同 (House Rental Contract)
    2. 借款合同 (Loan Agreement)
    3. 劳动合同 (Labor Contract)
    4. 法律咨询意见书 (Legal Opinion)
    5. 委托代理协议 (Agency Agreement)
    6. 保密协议 (NDA)
    """